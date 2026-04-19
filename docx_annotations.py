"""
docx_annotations.py - Erlaeuterungstexte aus Word-Dokumenten fuer EBA-DPM-Tabellenblaetter
============================================================================================

Durchsucht den data-Ordner (bzw. den Ordner der Excel-Datei) nach *.DOCX-Dateien,
die als Markdown-Text vorliegen (wie die FISMA-Quelldateien), und ordnet die
Textabschnitte den einzelnen Tabellenblaettern anhand der Meldeborgen-Codes zu.

Oeffentliche API:
    get_sheet_annotations(excel_path: str) -> dict[str, str]
        Gibt ein Dict zurueck: sheet_name -> HTML-String mit dem Erlaeuterungstext.
        Wenn keine DOCX-Dateien vorhanden sind, wird ein leeres Dict zurueckgegeben.

    render_annotation(annotation_html: str) -> None
        Rendert den Erlaeuterungstext in Streamlit als aufklappbaren Bereich.
"""

from __future__ import annotations

import re
from pathlib import Path


# ---------------------------------------------------------------------------
# Hilfsfunktionen zum Lesen der Dateien
# ---------------------------------------------------------------------------

def _read_file(path):
    try:
        return Path(path).read_text(encoding="utf-8")
    except Exception:
        try:
            return Path(path).read_text(encoding="latin-1")
        except Exception:
            return ""


def _is_text_file(path):
    """Prueft ob es eine Text-Datei ist (trotz .DOCX-Endung). Echte DOCX beginnen mit PK."""
    try:
        header = Path(path).read_bytes()[:4]
        return header[:2] != b"PK"
    except Exception:
        return False


def _try_read_real_docx(path):
    try:
        import docx
        doc = docx.Document(path)
        lines = [para.text for para in doc.paragraphs]
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        return "\n".join(lines)
    except Exception:
        return ""


# ---------------------------------------------------------------------------
# Parse-Logik
# ---------------------------------------------------------------------------

_SECTION_RE = re.compile(
    r"(?:^-\s+|\s*\d+\.\s+|^)(C\s*[\d]+(?:\.\d+)*)\s*[-\u2013\u2014]?\s*(.*)",
    re.IGNORECASE,
)


def _normalize_code(raw):
    return re.sub(r"\s+", "", raw.upper())


def _parse_sections(content):
    lines = content.split("\n")
    sections = {}
    current_key = None
    intro_lines = []
    found_section = False

    for line in lines:
        m = _SECTION_RE.match(line.strip())
        if m:
            found_section = True
            key = _normalize_code(m.group(1))
            title_rest = m.group(2).strip()
            current_key = key
            sections.setdefault(key, [])
            header = "### " + m.group(1).strip()
            if title_rest:
                header += " \u2013 " + title_rest
            sections[key].append(header)
        elif current_key is not None:
            sections[current_key].append(line)
        elif not found_section:
            intro_lines.append(line)

    result = {k: "\n".join(v).strip() for k, v in sections.items()}
    if intro_lines:
        result["__intro__"] = "\n".join(intro_lines).strip()
    return result


# ---------------------------------------------------------------------------
# Index-Mapping
# ---------------------------------------------------------------------------

def _build_code_to_sheet_map(excel_path):
    try:
        import openpyxl
    except ImportError:
        return {}
    try:
        wb = openpyxl.load_workbook(excel_path, data_only=True, read_only=True)
    except Exception:
        return {}

    if "Index" not in wb.sheetnames:
        wb.close()
        return {}

    mapping = {}
    ws = wb["Index"]
    for row in ws.iter_rows(min_row=1, max_row=600, values_only=True):
        if not row or len(row) < 5:
            continue
        raw_code = str(row[2]).strip() if row[2] else ""
        raw_sheet = str(row[4]).strip() if row[4] else ""
        if not raw_code or not raw_sheet or raw_code == "None" or raw_sheet == "None":
            continue
        norm = _normalize_code(raw_code)
        if not re.match(r"^C\d", norm):
            continue
        sheet = raw_sheet.replace("\xa0", " ").strip()
        mapping.setdefault(norm, [])
        if sheet not in mapping[norm]:
            mapping[norm].append(sheet)

    wb.close()
    return mapping


# ---------------------------------------------------------------------------
# Hauptfunktion
# ---------------------------------------------------------------------------

def get_sheet_annotations(excel_path):
    """
    Durchsucht den Ordner der Excel-Datei nach *.DOCX-Dateien.
    Gibt {sheet_name: html_text, ...} zurueck. Leer wenn keine Dateien gefunden.
    """
    data_dir = Path(excel_path).parent
    docx_files = sorted(
        list(data_dir.glob("*.DOCX")) + list(data_dir.glob("*.docx"))
    )

    if not docx_files:
        return {}

    code_to_sheets = _build_code_to_sheet_map(excel_path)

    all_sections = {}
    for fpath in docx_files:
        path_str = str(fpath)
        content = _read_file(path_str) if _is_text_file(path_str) else _try_read_real_docx(path_str)
        if not content.strip():
            continue
        sections = _parse_sections(content)
        sections.pop("__intro__", None)
        for code, text in sections.items():
            if code not in all_sections:
                all_sections[code] = text
            else:
                all_sections[code] += "\n\n" + text

    if not all_sections:
        return {}

    sheet_annotations = {}
    for norm_code, md_text in all_sections.items():
        sheets = code_to_sheets.get(norm_code, [])
        if not sheets:
            prefix_m = re.match(r"(C\d+)", norm_code)
            if prefix_m:
                p = prefix_m.group(1)
                sheets = [s for k, s_list in code_to_sheets.items() if k.startswith(p) for s in s_list]

        html = _markdown_to_html(md_text)
        for sheet in sheets:
            if sheet not in sheet_annotations:
                sheet_annotations[sheet] = html
            else:
                sheet_annotations[sheet] += html

    return sheet_annotations


# ---------------------------------------------------------------------------
# Markdown -> HTML
# ---------------------------------------------------------------------------

def _inline_md(text):
    text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<strong><em>\1</em></strong>", text)
    text = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", text)
    text = re.sub(r"\*(.+?)\*", r"<em>\1</em>", text)
    text = re.sub(r"`(.+?)`", r"<code>\1</code>", text)
    return text


def _markdown_to_html(md):
    lines = md.split("\n")
    html_parts = []
    in_ul = False
    in_table = False
    table_buffer = []

    def flush_table():
        nonlocal in_table, table_buffer
        if not table_buffer:
            in_table = False
            return
        rows = [r for r in table_buffer if r.strip()]
        html_parts.append('<table class="doc-annotation-table">')
        first_data = True
        for row in rows:
            if re.match(r"^\|[\s\-|]+\|?$", row.strip()):
                continue
            cells = [c.strip() for c in row.strip().strip("|").split("|")]
            tag = "th" if first_data else "td"
            first_data = False
            html_parts.append("  <tr>")
            for cell in cells:
                html_parts.append("    <{0}>{1}</{0}>".format(tag, _inline_md(cell)))
            html_parts.append("  </tr>")
        html_parts.append("</table>")
        in_table = False
        table_buffer.clear()

    def flush_ul():
        nonlocal in_ul
        if in_ul:
            html_parts.append("</ul>")
            in_ul = False

    for line in lines:
        stripped = line.strip()

        if stripped.startswith("|"):
            if not in_table:
                flush_ul()
                in_table = True
                table_buffer.clear()
            table_buffer.append(stripped)
            continue
        elif in_table:
            flush_table()

        h_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if h_match:
            flush_ul()
            level = len(h_match.group(1))
            html_parts.append(
                '<h{0} class="doc-annotation-h{0}">{1}</h{0}>'.format(
                    level, _inline_md(h_match.group(2))
                )
            )
            continue

        li_match = re.match(r"^[-*]\s+(.*)", stripped)
        if li_match:
            if not in_ul:
                html_parts.append("<ul>")
                in_ul = True
            html_parts.append("  <li>{0}</li>".format(_inline_md(li_match.group(1))))
            continue

        num_match = re.match(r"^-?\s*(\d+)\.\s+(.*)", stripped)
        if num_match:
            flush_ul()
            html_parts.append(
                '<p class="doc-annotation-numbered">'
                '<span class="num">{0}.</span> {1}</p>'.format(
                    num_match.group(1), _inline_md(num_match.group(2))
                )
            )
            continue

        flush_ul()
        if not stripped or stripped in ("(...)", "(\u2026)"):
            continue
        html_parts.append("<p>{0}</p>".format(_inline_md(stripped)))

    flush_ul()
    if in_table:
        flush_table()

    return "\n".join(html_parts)


# ---------------------------------------------------------------------------
# CSS (als String, ohne problematische Zeichen im Python-Quelltext)
# ---------------------------------------------------------------------------

_ANNOTATION_CSS = "\n".join([
    "<style>",
    ".doc-annotation-container {",
    "    background: #FFF8F8;",          # SR: sehr helles Cremeweiß
    "    border: 1px solid #DDDDDD;",
    "    border-left: 4px solid #CC0000;",  # SR: Sparkassen-Rot
    "    border-radius: 6px;",
    "    padding: 16px 20px 12px 20px;",
    "    margin-bottom: 16px;",
    "    font-size: 0.88rem;",
    "    line-height: 1.6;",
    "    color: #1A1A1A;",               # SR: Fast-Schwarz
    "}",
    ".doc-annotation-container h3.doc-annotation-h3 {",
    "    font-size: 1rem;",
    "    color: #1A1A1A;",
    "    margin: 10px 0 4px 0;",
    "    padding-bottom: 2px;",
    "    border-bottom: 1px solid #DDDDDD;",
    "}",
    ".doc-annotation-container h4.doc-annotation-h4 {",
    "    font-size: 0.92rem;",
    "    color: #990000;",               # SR: Dunkelrot vertieft
    "    margin: 8px 0 2px 0;",
    "}",
    ".doc-annotation-container p {",
    "    margin: 4px 0 6px 0;",
    "}",
    ".doc-annotation-container p.doc-annotation-numbered {",
    "    margin: 3px 0;",
    "    display: flex;",
    "    gap: 8px;",
    "}",
    ".doc-annotation-container p.doc-annotation-numbered .num {",
    "    color: #CC0000;",               # SR: Sparkassen-Rot
    "    font-weight: 600;",
    "    min-width: 24px;",
    "}",
    ".doc-annotation-container ul {",
    "    margin: 4px 0 6px 24px;",
    "    padding: 0;",
    "}",
    ".doc-annotation-container li {",
    "    margin: 2px 0;",
    "}",
    ".doc-annotation-container table.doc-annotation-table {",
    "    border-collapse: collapse;",
    "    width: 100%;",
    "    font-size: 0.83rem;",
    "    margin: 8px 0 12px 0;",
    "}",
    ".doc-annotation-container table.doc-annotation-table th,",
    ".doc-annotation-container table.doc-annotation-table td {",
    "    border: 1px solid #DDDDDD;",
    "    padding: 5px 10px;",
    "    vertical-align: top;",
    "    text-align: left;",
    "}",
    ".doc-annotation-container table.doc-annotation-table th {",
    "    background: #FDECEA;",          # SR: sehr helles Rot
    "    font-weight: 600;",
    "    color: #990000;",               # SR: Dunkelrot vertieft
    "}",
    ".doc-annotation-container table.doc-annotation-table tr:nth-child(even) td {",
    "    background: #FFF8F8;",          # SR: sehr helles Cremeweiß
    "}",
    "</style>",
])


# ---------------------------------------------------------------------------
# Streamlit-Render-Funktion
# ---------------------------------------------------------------------------

def render_annotation(annotation_html):
    """
    Rendert den Erlaeuterungstext als aufklappbaren Expander vor der Tabelle.
    Wenn kein Text vorhanden ist, wird nichts ausgegeben (Status-Quo bleibt).
    """
    if not annotation_html:
        return

    import streamlit as st

    if "_doc_annotation_css_injected" not in st.session_state:
        st.markdown(_ANNOTATION_CSS, unsafe_allow_html=True)
        st.session_state["_doc_annotation_css_injected"] = True

    with st.expander("\U0001f4c4 Erlaeuterungen (Rechtsgrundlagen)", expanded=False):
        st.markdown(
            '<div class="doc-annotation-container">' + annotation_html + "</div>",
            unsafe_allow_html=True,
        )
